from __future__ import annotations

import zipfile
from collections.abc import Callable, Iterator, Mapping, Sequence
from datetime import UTC, datetime, timedelta
from io import BytesIO
from pathlib import Path
from typing import cast

import pytest
from docx import Document
from sqlalchemy import create_engine, select
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

from app.core.config import get_settings
from app.core.job_utils import LeaseLostError
from app.core.security import hash_password
from app.db.base import Base
from app.db.models import DocumentChunk, DocumentVersion, Job, LogicalDocument, Role, User
from app.ingest.embedding import (
    DocumentEmbeddingService,
    EmbeddingAdapterError,
    EmbeddingBatchConfig,
    FakeEmbeddingAdapter,
)
from app.ingest.extractors.base import (
    ExtractedDocument,
    ExtractedPage,
    ExtractionInputMetadata,
    ExtractionMetadata,
)
from app.ingest.extractors.dispatcher import ExtractorDispatcher
from app.ingest.qdrant import (
    DocumentIndexingService,
    InMemoryQdrantClient,
    QdrantCollectionConfig,
    QdrantPoint,
    QdrantStoreError,
    QdrantVectorStore,
)
from app.repositories.document_repository import DocumentRepository
from app.repositories.job_repository import JobRepository
from app.storage.file_storage import LocalFileStorage
from app.workers.handlers.base import JobExecutionContext
from app.workers.handlers.document_ingest_handler import DocumentIngestHandler
from app.workers.job_dispatcher import JobDispatcher
from app.workers.worker_config import WorkerConfig
from app.workers.worker_main import WorkerRunner

TEST_PASSWORD = "password"


@pytest.fixture
def ingest_session_factory(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> Iterator[tuple[sessionmaker[Session], LocalFileStorage]]:
    monkeypatch.setenv("STORAGE_ROOT", str(tmp_path))
    monkeypatch.setenv(
        "UPLOAD_ALLOWED_EXTENSIONS",
        '[".pdf",".docx",".txt",".md",".markdown",".csv",".xlsx",".pptx"]',
    )
    monkeypatch.setenv("INGEST_CHUNK_SIZE_TOKENS", "5")
    monkeypatch.setenv("INGEST_CHUNK_OVERLAP_TOKENS", "1")
    monkeypatch.setenv("EMBEDDING_PROVIDER", "fake")
    monkeypatch.setenv("EMBEDDING_FAKE_DIMENSION", "4")
    monkeypatch.setenv("QDRANT_COLLECTION_NAME", "test_document_chunks")
    get_settings.cache_clear()
    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(engine)
    factory = sessionmaker(bind=engine, autoflush=False, autocommit=False)
    with factory() as db:
        admin_role = Role(role_name="admin", description="Admin")
        db.add(admin_role)
        db.flush()
        db.add(
            User(
                role_id=admin_role.role_id,
                email="admin@example.com",
                display_name="Admin",
                password_hash=hash_password(TEST_PASSWORD),
                status="active",
            )
        )
        db.commit()
    try:
        yield factory, LocalFileStorage(tmp_path)
    finally:
        get_settings.cache_clear()
        engine.dispose()


@pytest.mark.parametrize(
    ("file_name", "mime_type", "content"),
    [
        ("doc.txt", "text/plain", b"alpha beta gamma delta epsilon zeta"),
        ("doc.md", "text/markdown", b"# Intro\nalpha beta gamma\n## Next\ndelta epsilon zeta"),
        ("doc.csv", "text/csv", b"name,value\nalpha,1\nbeta,2\n"),
        ("doc.pdf", "application/pdf", None),
        (
            "doc.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            None,
        ),
    ],
)
def test_document_ingest_handler_success_for_supported_types(
    ingest_session_factory: tuple[sessionmaker[Session], LocalFileStorage],
    tmp_path: Path,
    file_name: str,
    mime_type: str,
    content: bytes | None,
) -> None:
    session_factory, storage = ingest_session_factory
    if file_name.endswith(".docx"):
        content = _docx_bytes(tmp_path)
    if file_name.endswith(".pdf"):
        content = _minimal_pdf("alpha beta gamma delta")
    assert content is not None
    version_id = _create_document_version(
        session_factory,
        storage,
        file_name=file_name,
        mime_type=mime_type,
        content=content,
    )

    result = _handler(session_factory, storage).handle(_context(version_id))

    assert result.status == "succeeded"
    with session_factory() as db:
        version = db.get(DocumentVersion, version_id)
        chunks = db.scalars(
            select(DocumentChunk)
            .where(DocumentChunk.document_version_id == version_id)
            .order_by(DocumentChunk.chunk_index)
        ).all()
        assert version is not None
        assert version.status == "ready"
        assert version.error_code is None
        assert version.extractor_name is not None
        assert chunks
        assert [chunk.chunk_index for chunk in chunks] == list(range(len(chunks)))
        assert all(chunk.chunk_hash and len(chunk.chunk_hash) == 64 for chunk in chunks)


@pytest.mark.parametrize(
    ("file_name", "mime_type", "content", "structure_type", "expected_label"),
    [
        (
            "sales.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            lambda: _minimal_xlsx(),
            "excel_sheet",
            "Sales",
        ),
        (
            "deck.pptx",
            "application/vnd.openxmlformats-officedocument.presentationml.presentation",
            lambda: _minimal_pptx(),
            "powerpoint_slide",
            "Architecture",
        ),
    ],
)
def test_document_ingest_handler_success_for_office_parent_child_chunks(
    ingest_session_factory: tuple[sessionmaker[Session], LocalFileStorage],
    file_name: str,
    mime_type: str,
    content: Callable[[], bytes],
    structure_type: str,
    expected_label: str,
) -> None:
    session_factory, storage = ingest_session_factory
    content_bytes = content()
    version_id = _create_document_version(
        session_factory,
        storage,
        file_name=file_name,
        mime_type=mime_type,
        content=content_bytes,
    )
    qdrant_client = InMemoryQdrantClient()

    result = _handler(
        session_factory,
        storage,
        indexing_service=_indexing_service(qdrant_client=qdrant_client),
    ).handle(_context(version_id))

    assert result.status == "succeeded"
    with session_factory() as db:
        version = db.get(DocumentVersion, version_id)
        chunks = db.scalars(
            select(DocumentChunk)
            .where(DocumentChunk.document_version_id == version_id)
            .order_by(DocumentChunk.chunk_index)
        ).all()
        assert version is not None
        assert version.status == "ready"
        assert version.extractor_name in {"xlsx", "pptx"}
        assert chunks
        metadata = chunks[0].metadata_json
        assert metadata is not None
        assert metadata["parent_child_schema_version"] == "phase2.parent_child.v1"
        assert metadata["structure_type"] == structure_type
        assert metadata["chunk_level"] == "child"
        assert expected_label in (chunks[0].section_title or "")
        points = qdrant_client.points["test_document_chunks"]
        assert points
        first_payload = next(iter(points.values())).payload
        assert first_payload["structure_type"] == structure_type
        assert "content_text" not in first_payload
        assert "raw_chunk_text" not in first_payload
        if structure_type == "excel_sheet":
            assert first_payload["sheet_name"] == "Sales"
        else:
            assert first_payload["slide_number"] == 1


def test_document_ingest_handler_missing_storage_file_fails_safely(
    ingest_session_factory: tuple[sessionmaker[Session], LocalFileStorage],
) -> None:
    session_factory, storage = ingest_session_factory
    version_id = _create_document_version(
        session_factory,
        storage,
        file_name="missing.txt",
        mime_type="text/plain",
        content=None,
    )
    _set_stale_ingest_metadata(session_factory, version_id)

    result = _handler(session_factory, storage).handle(_context(version_id))

    assert result.status == "failed"
    assert result.error_code == "storage_file_missing"
    assert result.error_message == "Stored object was not found."
    _assert_failed_version(session_factory, version_id, "storage_file_missing")


def test_document_ingest_handler_rejects_invalid_payload_before_domain_update(
    ingest_session_factory: tuple[sessionmaker[Session], LocalFileStorage],
) -> None:
    session_factory, storage = ingest_session_factory
    version_id = _create_document_version(
        session_factory,
        storage,
        file_name="valid.txt",
        mime_type="text/plain",
        content=b"alpha beta",
    )

    invalid_bool = _handler(session_factory, storage).handle(
        JobExecutionContext(
            job_id=1,
            job_type="document_ingest",
            target_type="document_version",
            target_id=version_id,
            payload={"document_version_id": True},
            worker_instance_id="worker-1",
        )
    )
    mismatch = _handler(session_factory, storage).handle(
        JobExecutionContext(
            job_id=2,
            job_type="document_ingest",
            target_type="document_version",
            target_id=version_id + 1,
            payload={"document_version_id": version_id},
            worker_instance_id="worker-1",
        )
    )
    logical_document_mismatch = _handler(session_factory, storage).handle(
        JobExecutionContext(
            job_id=3,
            job_type="document_ingest",
            target_type="document_version",
            target_id=version_id,
            payload={"document_version_id": version_id, "logical_document_id": 999},
            worker_instance_id="worker-1",
        )
    )
    invalid_logical_id_bool = _handler(session_factory, storage).handle(
        JobExecutionContext(
            job_id=4,
            job_type="document_ingest",
            target_type="document_version",
            target_id=version_id,
            payload={"document_version_id": version_id, "logical_document_id": True},
            worker_instance_id="worker-1",
        )
    )

    assert invalid_bool.status == "failed"
    assert invalid_bool.error_code == "validation_error"
    assert mismatch.status == "failed"
    assert mismatch.error_code == "validation_error"
    assert logical_document_mismatch.status == "failed"
    assert logical_document_mismatch.error_code == "validation_error"
    assert invalid_logical_id_bool.status == "failed"
    assert invalid_logical_id_bool.error_code == "validation_error"
    with session_factory() as db:
        version = db.get(DocumentVersion, version_id)
        assert version is not None
        assert version.status == "processing"
        assert version.error_code is None


def test_document_ingest_handler_ready_version_is_noop_and_preserves_chunks(
    ingest_session_factory: tuple[sessionmaker[Session], LocalFileStorage],
) -> None:
    session_factory, storage = ingest_session_factory
    version_id = _create_document_version(
        session_factory,
        storage,
        file_name="ready.txt",
        mime_type="text/plain",
        content=b"ready alpha beta",
        status="ready",
    )
    with session_factory() as db:
        db.add(
            DocumentChunk(
                document_version_id=version_id,
                chunk_index=0,
                chunk_hash="c" * 64,
                content_text="existing chunk",
            )
        )
        db.commit()

    result = _handler(session_factory, storage).handle(_context(version_id))

    assert result.status == "succeeded"
    assert result.result_json["result_code"] == "no_op"
    with session_factory() as db:
        version = db.get(DocumentVersion, version_id)
        assert version is not None
        assert version.status == "ready"
        chunks = db.scalars(select(DocumentChunk)).all()
        assert len(chunks) == 1
        assert chunks[0].content_text == "existing chunk"


def test_document_ingest_handler_unsupported_and_empty_text_fail_safely(
    ingest_session_factory: tuple[sessionmaker[Session], LocalFileStorage],
) -> None:
    session_factory, storage = ingest_session_factory
    unsupported_id = _create_document_version(
        session_factory,
        storage,
        file_name="script.exe",
        mime_type="application/octet-stream",
        content=b"MZ",
    )
    empty_id = _create_document_version(
        session_factory,
        storage,
        file_name="empty.txt",
        mime_type="text/plain",
        content=b"   \n",
    )

    unsupported = _handler(session_factory, storage).handle(_context(unsupported_id))
    empty = _handler(session_factory, storage).handle(_context(empty_id))

    assert unsupported.status == "failed"
    assert unsupported.error_code == "unsupported_file_type"
    assert empty.status == "failed"
    assert empty.error_code == "empty_extracted_text"
    _assert_failed_version(session_factory, unsupported_id, "unsupported_file_type")
    _assert_failed_version(session_factory, empty_id, "empty_extracted_text")


def test_document_ingest_handler_chunking_zero_failure_cleans_partial_chunks(
    ingest_session_factory: tuple[sessionmaker[Session], LocalFileStorage],
) -> None:
    session_factory, storage = ingest_session_factory
    version_id = _create_document_version(
        session_factory,
        storage,
        file_name="blank.txt",
        mime_type="text/plain",
        content=b"ignored",
    )
    with session_factory() as db:
        db.add(
            DocumentChunk(
                document_version_id=version_id,
                chunk_index=0,
                chunk_hash="a" * 64,
                content_text="old",
            )
        )
        db.commit()

    handler = DocumentIngestHandler(
        session_factory=session_factory,
        job_repository=cast(JobRepository, _NoopJobRepository()),
        storage=storage,
        dispatcher=cast(ExtractorDispatcher, _WhitespaceDispatcher()),
        settings=get_settings(),
        indexing_service=_indexing_service(),
    )
    result = handler.handle(_context(version_id))

    assert result.status == "failed"
    assert result.error_code == "no_chunks_created"
    with session_factory() as db:
        assert db.query(DocumentChunk).filter_by(document_version_id=version_id).count() == 0
    _assert_failed_version(session_factory, version_id, "no_chunks_created")


def test_document_ingest_handler_insert_failure_marks_failed_and_cleans_chunks(
    ingest_session_factory: tuple[sessionmaker[Session], LocalFileStorage],
) -> None:
    session_factory, storage = ingest_session_factory
    version_id = _create_document_version(
        session_factory,
        storage,
        file_name="insert-failure.txt",
        mime_type="text/plain",
        content=b"alpha beta gamma delta",
    )
    with session_factory() as db:
        db.add(
            DocumentChunk(
                document_version_id=version_id,
                chunk_index=0,
                chunk_hash="d" * 64,
                content_text="old chunk",
            )
        )
        db.commit()

    handler = DocumentIngestHandler(
        session_factory=session_factory,
        repository=_FailingInsertDocumentRepository(),
        job_repository=cast(JobRepository, _NoopJobRepository()),
        storage=storage,
        settings=get_settings(),
        indexing_service=_indexing_service(),
    )
    result = handler.handle(_context(version_id))

    assert result.status == "failed"
    assert result.error_code == "document_chunk_insert_failed"
    _assert_failed_version(session_factory, version_id, "document_chunk_insert_failed")


def test_document_ingest_handler_extracted_text_over_limit_fails_safely(
    ingest_session_factory: tuple[sessionmaker[Session], LocalFileStorage],
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    session_factory, storage = ingest_session_factory
    monkeypatch.setenv("INGEST_MAX_EXTRACTED_TEXT_CHARS", "3")
    get_settings.cache_clear()
    version_id = _create_document_version(
        session_factory,
        storage,
        file_name="too-long.txt",
        mime_type="text/plain",
        content=b"alpha beta",
    )

    result = _handler(session_factory, storage).handle(_context(version_id))

    assert result.status == "failed"
    assert result.error_code == "text_extraction_failed"
    _assert_failed_version(session_factory, version_id, "text_extraction_failed")


def test_document_ingest_handler_lease_lost_blocks_domain_write(
    ingest_session_factory: tuple[sessionmaker[Session], LocalFileStorage],
) -> None:
    session_factory, storage = ingest_session_factory
    version_id = _create_document_version(
        session_factory,
        storage,
        file_name="lease.txt",
        mime_type="text/plain",
        content=b"alpha beta",
    )
    handler = DocumentIngestHandler(
        session_factory=session_factory,
        job_repository=cast(JobRepository, _LeaseLostJobRepository()),
        storage=storage,
        settings=get_settings(),
        indexing_service=_indexing_service(),
    )

    with pytest.raises(LeaseLostError):
        handler.handle(_context(version_id))

    with session_factory() as db:
        version = db.get(DocumentVersion, version_id)
        assert version is not None
        assert version.status == "processing"
        assert version.error_code is None
        assert db.query(DocumentChunk).filter_by(document_version_id=version_id).count() == 0


def test_document_ingest_handler_rechecks_archived_document_before_final_write(
    ingest_session_factory: tuple[sessionmaker[Session], LocalFileStorage],
) -> None:
    session_factory, storage = ingest_session_factory
    version_id = _create_document_version(
        session_factory,
        storage,
        file_name="archive-race.txt",
        mime_type="text/plain",
        content=b"alpha beta",
    )
    handler = DocumentIngestHandler(
        session_factory=session_factory,
        job_repository=cast(JobRepository, _NoopJobRepository()),
        storage=storage,
        dispatcher=cast(ExtractorDispatcher, _ArchivingDispatcher(session_factory)),
        settings=get_settings(),
        indexing_service=_indexing_service(),
    )

    result = handler.handle(_context(version_id))

    assert result.status == "failed"
    assert result.error_code == "document_version_not_ingestable"
    _assert_failed_version(session_factory, version_id, "document_version_not_ingestable")


def test_document_ingest_handler_retry_cleans_existing_chunks_and_reinserts(
    ingest_session_factory: tuple[sessionmaker[Session], LocalFileStorage],
) -> None:
    session_factory, storage = ingest_session_factory
    version_id = _create_document_version(
        session_factory,
        storage,
        file_name="retry.txt",
        mime_type="text/plain",
        content=b"new alpha beta gamma delta epsilon",
        status="failed",
        error_code="text_extraction_failed",
    )
    with session_factory() as db:
        db.add(
            DocumentChunk(
                document_version_id=version_id,
                chunk_index=0,
                chunk_hash="b" * 64,
                content_text="old chunk",
            )
        )
        db.commit()

    result = _handler(session_factory, storage).handle(_context(version_id))

    assert result.status == "succeeded"
    assert result.result_json["status"] == "ready"
    assert result.result_json["indexed_count"] == result.result_json["chunk_count"]
    with session_factory() as db:
        version = db.get(DocumentVersion, version_id)
        chunks = db.scalars(
            select(DocumentChunk).where(DocumentChunk.document_version_id == version_id)
        ).all()
        assert version is not None
        assert version.status == "ready"
        assert version.error_code is None
        assert chunks
        assert all(chunk.content_text != "old chunk" for chunk in chunks)


def test_worker_single_iteration_processes_document_ingest_success_and_failure(
    ingest_session_factory: tuple[sessionmaker[Session], LocalFileStorage],
) -> None:
    session_factory, storage = ingest_session_factory
    success_version_id = _create_document_version(
        session_factory,
        storage,
        file_name="worker.txt",
        mime_type="text/plain",
        content=b"worker alpha beta gamma delta epsilon",
    )
    failure_version_id = _create_document_version(
        session_factory,
        storage,
        file_name="worker-missing.txt",
        mime_type="text/plain",
        content=None,
    )
    with session_factory() as db:
        db.add_all(
            [
                Job(
                    job_id=1,
                    job_type="document_ingest",
                    status="queued",
                    target_type="document_version",
                    target_id=success_version_id,
                    payload_json={"document_version_id": success_version_id},
                ),
                Job(
                    job_id=2,
                    job_type="document_ingest",
                    status="queued",
                    target_type="document_version",
                    target_id=failure_version_id,
                    payload_json={"document_version_id": failure_version_id},
                ),
            ]
        )
        db.commit()

    dispatcher = JobDispatcher(
        {"document_ingest": _handler(session_factory, storage, enforce_lease=True)}
    )
    runner = WorkerRunner(
        config=_worker_config(batch_size=2),
        session_factory=session_factory,
        dispatcher=dispatcher,
    )

    assert runner.run_once() == 2

    with session_factory() as db:
        jobs = {job.job_id: job for job in db.scalars(select(Job)).all()}
        assert jobs[1].status == "succeeded"
        assert jobs[1].result_json == {
            "document_version_id": success_version_id,
            "logical_document_id": 1,
            "chunk_count": 2,
            "indexed_count": 2,
            "status": "ready",
        }
        assert jobs[2].status == "failed"
        assert jobs[2].error_code == "storage_file_missing"
        assert db.query(DocumentChunk).filter_by(document_version_id=success_version_id).count() > 0
        failed_version = db.get(DocumentVersion, failure_version_id)
        assert failed_version is not None
        assert failed_version.status == "failed"


def test_document_ingest_handler_embedding_failure_marks_failed_and_cleans_chunks(
    ingest_session_factory: tuple[sessionmaker[Session], LocalFileStorage],
) -> None:
    session_factory, storage = ingest_session_factory
    version_id = _create_document_version(
        session_factory,
        storage,
        file_name="embedding-failure.txt",
        mime_type="text/plain",
        content=b"alpha beta gamma delta",
    )
    handler = DocumentIngestHandler(
        session_factory=session_factory,
        job_repository=cast(JobRepository, _NoopJobRepository()),
        storage=storage,
        settings=get_settings(),
        indexing_service=_indexing_service(embedding_service=_FailingEmbeddingService()),
    )

    result = handler.handle(_context(version_id))

    assert result.status == "failed"
    assert result.error_code == "embedding_failed"
    _assert_failed_version(session_factory, version_id, "embedding_failed")


def test_document_ingest_handler_embedding_dimension_mismatch_marks_failed(
    ingest_session_factory: tuple[sessionmaker[Session], LocalFileStorage],
) -> None:
    session_factory, storage = ingest_session_factory
    version_id = _create_document_version(
        session_factory,
        storage,
        file_name="embedding-dimension.txt",
        mime_type="text/plain",
        content=b"alpha beta gamma delta",
    )
    handler = DocumentIngestHandler(
        session_factory=session_factory,
        job_repository=cast(JobRepository, _NoopJobRepository()),
        storage=storage,
        settings=get_settings(),
        indexing_service=_indexing_service(embedding_service=_DimensionMismatchEmbeddingService()),
    )

    result = handler.handle(_context(version_id))

    assert result.status == "failed"
    assert result.error_code == "embedding_dimension_mismatch"
    _assert_failed_version(session_factory, version_id, "embedding_dimension_mismatch")


def test_document_ingest_handler_qdrant_failure_marks_failed_and_cleans_partial_points(
    ingest_session_factory: tuple[sessionmaker[Session], LocalFileStorage],
) -> None:
    session_factory, storage = ingest_session_factory
    version_id = _create_document_version(
        session_factory,
        storage,
        file_name="qdrant-failure.txt",
        mime_type="text/plain",
        content=b"alpha beta gamma delta",
    )
    qdrant_client = InMemoryQdrantClient()
    qdrant_client.fail_upsert = True
    handler = DocumentIngestHandler(
        session_factory=session_factory,
        job_repository=cast(JobRepository, _NoopJobRepository()),
        storage=storage,
        settings=get_settings(),
        indexing_service=_indexing_service(qdrant_client=qdrant_client),
    )

    result = handler.handle(_context(version_id))

    assert result.status == "failed"
    assert result.error_code == "qdrant_upsert_failed"
    _assert_failed_version(session_factory, version_id, "qdrant_upsert_failed")


def test_document_ingest_handler_cleans_points_after_partial_batch_upsert_failure(
    ingest_session_factory: tuple[sessionmaker[Session], LocalFileStorage],
) -> None:
    session_factory, storage = ingest_session_factory
    version_id = _create_document_version(
        session_factory,
        storage,
        file_name="qdrant-partial.txt",
        mime_type="text/plain",
        content=b"alpha beta gamma delta epsilon zeta eta theta",
    )
    qdrant_client = _PartialFailingQdrantClient()
    handler = DocumentIngestHandler(
        session_factory=session_factory,
        job_repository=cast(JobRepository, _NoopJobRepository()),
        storage=storage,
        settings=get_settings(),
        indexing_service=_indexing_service(
            qdrant_client=qdrant_client,
            upsert_batch_size=1,
        ),
    )

    result = handler.handle(_context(version_id))

    assert result.status == "failed"
    assert result.error_code == "qdrant_upsert_failed"
    _assert_failed_version(session_factory, version_id, "qdrant_upsert_failed")
    assert qdrant_client.points["test_document_chunks"] == {}


def test_document_ingest_handler_ready_update_failure_attempts_qdrant_cleanup(
    ingest_session_factory: tuple[sessionmaker[Session], LocalFileStorage],
) -> None:
    session_factory, storage = ingest_session_factory
    version_id = _create_document_version(
        session_factory,
        storage,
        file_name="ready-update-failure.txt",
        mime_type="text/plain",
        content=b"alpha beta gamma delta",
    )
    qdrant_client = InMemoryQdrantClient()
    handler = DocumentIngestHandler(
        session_factory=session_factory,
        repository=_FailingReadyDocumentRepository(),
        job_repository=cast(JobRepository, _NoopJobRepository()),
        storage=storage,
        settings=get_settings(),
        indexing_service=_indexing_service(qdrant_client=qdrant_client),
    )

    result = handler.handle(_context(version_id))

    assert result.status == "failed"
    assert result.error_code == "document_ready_update_failed"
    _assert_failed_version(session_factory, version_id, "document_ready_update_failed")
    assert qdrant_client.points["test_document_chunks"] == {}


def test_document_ingest_handler_preserves_chunks_when_qdrant_cleanup_fails(
    ingest_session_factory: tuple[sessionmaker[Session], LocalFileStorage],
) -> None:
    session_factory, storage = ingest_session_factory
    version_id = _create_document_version(
        session_factory,
        storage,
        file_name="ready-update-cleanup-failure.txt",
        mime_type="text/plain",
        content=b"alpha beta gamma delta",
    )
    qdrant_client = InMemoryQdrantClient()
    qdrant_client.fail_delete = True
    handler = DocumentIngestHandler(
        session_factory=session_factory,
        repository=_FailingReadyDocumentRepository(),
        job_repository=cast(JobRepository, _NoopJobRepository()),
        storage=storage,
        settings=get_settings(),
        indexing_service=_indexing_service(qdrant_client=qdrant_client),
    )

    result = handler.handle(_context(version_id))

    assert result.status == "failed"
    assert result.error_code == "document_ready_update_failed"
    with session_factory() as db:
        version = db.get(DocumentVersion, version_id)
        chunks = db.scalars(
            select(DocumentChunk).where(DocumentChunk.document_version_id == version_id)
        ).all()
        assert version is not None
        assert version.status == "failed"
        assert version.error_code == "document_ready_update_failed"
        assert chunks


def test_document_ingest_handler_retry_cleans_existing_qdrant_points_before_reinsert(
    ingest_session_factory: tuple[sessionmaker[Session], LocalFileStorage],
) -> None:
    session_factory, storage = ingest_session_factory
    version_id = _create_document_version(
        session_factory,
        storage,
        file_name="retry-index.txt",
        mime_type="text/plain",
        content=b"new alpha beta gamma delta epsilon",
        status="failed",
        error_code="qdrant_upsert_failed",
    )
    qdrant_client = InMemoryQdrantClient()
    with session_factory() as db:
        db.add(
            DocumentChunk(
                document_chunk_id=99,
                document_version_id=version_id,
                chunk_index=0,
                chunk_hash="e" * 64,
                content_text="old chunk",
            )
        )
        db.commit()
    qdrant_client.create_collection(
        QdrantCollectionConfig(
            name="test_document_chunks",
            vector_dimension=get_settings().effective_embedding_dimension,
        )
    )
    qdrant_client.upsert_points(
        "test_document_chunks",
        [
            QdrantPoint(
                point_id=99,
                vector=[0.0] * get_settings().effective_embedding_dimension,
                payload={"document_version_id": version_id, "document_chunk_id": 99},
            )
        ],
    )
    handler = DocumentIngestHandler(
        session_factory=session_factory,
        job_repository=cast(JobRepository, _NoopJobRepository()),
        storage=storage,
        settings=get_settings(),
        indexing_service=_indexing_service(qdrant_client=qdrant_client),
    )

    result = handler.handle(_context(version_id))

    assert result.status == "succeeded"
    with session_factory() as db:
        chunks = db.scalars(
            select(DocumentChunk).where(DocumentChunk.document_version_id == version_id)
        ).all()
        assert chunks
        assert all(chunk.document_chunk_id != 99 for chunk in chunks)
    assert 99 not in qdrant_client.points["test_document_chunks"]


class _WhitespaceExtractor:
    name = "whitespace"
    version = "1"

    def extract(self, file_path: Path, metadata: ExtractionInputMetadata) -> ExtractedDocument:
        return ExtractedDocument(
            pages=[ExtractedPage("   ", page_number=None)],
            metadata=ExtractionMetadata(
                extractor_name=self.name,
                extractor_version=self.version,
                page_count=None,
            ),
        )


class _WhitespaceDispatcher:
    def select(self, *, file_name: str, mime_type: str) -> _WhitespaceExtractor:
        return _WhitespaceExtractor()


class _ArchivingExtractor:
    name = "archiving"
    version = "1"

    def __init__(self, session_factory: sessionmaker[Session]) -> None:
        self.session_factory = session_factory

    def extract(self, file_path: Path, metadata: ExtractionInputMetadata) -> ExtractedDocument:
        with self.session_factory() as db:
            document = db.get(LogicalDocument, 1)
            assert document is not None
            document.status = "archived"
            document.archived_at = datetime.now(UTC)
            db.commit()
        return ExtractedDocument(
            pages=[ExtractedPage("alpha beta", page_number=None)],
            metadata=ExtractionMetadata(
                extractor_name=self.name,
                extractor_version=self.version,
                page_count=None,
            ),
        )


class _ArchivingDispatcher:
    def __init__(self, session_factory: sessionmaker[Session]) -> None:
        self.session_factory = session_factory

    def select(self, *, file_name: str, mime_type: str) -> _ArchivingExtractor:
        return _ArchivingExtractor(self.session_factory)


def _handler(
    session_factory: sessionmaker[Session],
    storage: LocalFileStorage,
    *,
    enforce_lease: bool = False,
    indexing_service: DocumentIndexingService | None = None,
) -> DocumentIngestHandler:
    return DocumentIngestHandler(
        session_factory=session_factory,
        job_repository=None if enforce_lease else cast(JobRepository, _NoopJobRepository()),
        storage=storage,
        settings=get_settings(),
        indexing_service=indexing_service or _indexing_service(),
    )


def _context(document_version_id: int) -> JobExecutionContext:
    return JobExecutionContext(
        job_id=100 + document_version_id,
        job_type="document_ingest",
        target_type="document_version",
        target_id=document_version_id,
        payload={"document_version_id": document_version_id, "logical_document_id": 1},
        worker_instance_id="worker-1",
    )


def _create_document_version(
    session_factory: sessionmaker[Session],
    storage: LocalFileStorage,
    *,
    file_name: str,
    mime_type: str,
    content: bytes | None,
    status: str = "processing",
    error_code: str | None = None,
) -> int:
    storage_key = storage.build_storage_key(file_name=file_name)
    if content is not None:
        storage.save_bytes(storage_key=storage_key, content=content)
    with session_factory() as db:
        user = db.scalar(select(User).where(User.email == "admin@example.com"))
        assert user is not None
        document = db.scalar(
            select(LogicalDocument).where(LogicalDocument.logical_document_id == 1)
        )
        if document is None:
            document = LogicalDocument(
                logical_document_id=1,
                owner_user_id=user.user_id,
                title="Ingest",
            )
            db.add(document)
            db.flush()
        version_no = db.query(DocumentVersion).count() + 1
        version = DocumentVersion(
            logical_document_id=document.logical_document_id,
            version_no=version_no,
            content_hash=f"{version_no:064x}",
            status=status,
            error_code=error_code,
            file_name=file_name,
            mime_type=mime_type,
            file_size_bytes=len(content or b""),
            storage_key=storage_key,
            created_by=user.user_id,
        )
        db.add(version)
        db.commit()
        return int(version.document_version_id)


def _assert_failed_version(
    session_factory: sessionmaker[Session], document_version_id: int, error_code: str
) -> None:
    with session_factory() as db:
        version = db.get(DocumentVersion, document_version_id)
        assert version is not None
        assert version.status == "failed"
        assert version.error_code == error_code
        assert version.page_count is None
        assert version.extractor_name is None
        assert version.extractor_version is None
        assert (
            db.query(DocumentChunk).filter_by(document_version_id=document_version_id).count() == 0
        )


def _set_stale_ingest_metadata(
    session_factory: sessionmaker[Session], document_version_id: int
) -> None:
    with session_factory() as db:
        version = db.get(DocumentVersion, document_version_id)
        assert version is not None
        version.page_count = 99
        version.extractor_name = "stale"
        version.extractor_version = "stale"
        db.commit()


class _NoopJobRepository:
    def assert_ownership(self, db: Session, *, job_id: int, worker_instance_id: str) -> None:
        return None


class _LeaseLostJobRepository:
    def assert_ownership(self, db: Session, *, job_id: int, worker_instance_id: str) -> None:
        raise LeaseLostError(f"Lease lost for job_id={job_id}")


class _FailingInsertDocumentRepository(DocumentRepository):
    def bulk_insert_chunks(
        self,
        db: Session,
        *,
        chunks: Sequence[Mapping[str, object]],
    ) -> None:
        raise RuntimeError("safe synthetic insert failure")


class _FailingReadyDocumentRepository(DocumentRepository):
    def mark_version_ready(
        self,
        db: Session,
        *,
        version: DocumentVersion,
        updated_at: datetime,
    ) -> None:
        raise RuntimeError("safe synthetic ready update failure")


class _FailingEmbeddingService:
    def embed_chunks(self, chunks: Sequence[object]) -> list[list[float]]:
        raise RuntimeError("safe synthetic embedding failure")


class _DimensionMismatchEmbeddingService:
    def embed_chunks(self, chunks: Sequence[object]) -> list[list[float]]:
        raise EmbeddingAdapterError("embedding_dimension_mismatch")


class _PartialFailingQdrantClient(InMemoryQdrantClient):
    def __init__(self) -> None:
        super().__init__()
        self.upsert_calls = 0

    def upsert_points(self, collection_name: str, points: Sequence[QdrantPoint]) -> None:
        self.upsert_calls += 1
        if self.upsert_calls == 1:
            super().upsert_points(collection_name, points)
            return
        raise QdrantStoreError("qdrant_upsert_failed")


def _indexing_service(
    *,
    qdrant_client: InMemoryQdrantClient | None = None,
    embedding_service: (
        DocumentEmbeddingService
        | _FailingEmbeddingService
        | _DimensionMismatchEmbeddingService
        | None
    ) = None,
    upsert_batch_size: int | None = None,
) -> DocumentIndexingService:
    settings = get_settings()
    dimension = settings.effective_embedding_dimension
    effective_embedding_service = embedding_service or DocumentEmbeddingService(
        adapter=FakeEmbeddingAdapter(dimension=dimension),
        config=EmbeddingBatchConfig(dimension=dimension, batch_size=settings.embedding_batch_size),
    )
    return DocumentIndexingService(
        embedding_service=cast(DocumentEmbeddingService, effective_embedding_service),
        vector_store=QdrantVectorStore(
            client=qdrant_client or InMemoryQdrantClient(),
            config=QdrantCollectionConfig(
                name=settings.qdrant_collection_name,
                vector_dimension=dimension,
                distance=settings.qdrant_distance,
            ),
            create_collection=True,
        ),
        upsert_batch_size=upsert_batch_size or settings.qdrant_upsert_batch_size,
    )


def _docx_bytes(tmp_path: Path) -> bytes:
    path = tmp_path / "fixture.docx"
    document = Document()
    document.add_paragraph("alpha beta gamma delta")
    document.save(str(path))
    return path.read_bytes()


def _minimal_pdf(text: str) -> bytes:
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("ascii")
    objects = [
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n",
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n",
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >> endobj\n",
        b"4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n",
        b"5 0 obj << /Length "
        + str(len(stream)).encode("ascii")
        + b" >> stream\n"
        + stream
        + b"\nendstream endobj\n",
    ]
    content = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for obj in objects:
        offsets.append(len(content))
        content.extend(obj)
    xref_offset = len(content)
    content.extend(f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode("ascii"))
    for offset in offsets[1:]:
        content.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    content.extend(
        (
            f"trailer << /Root 1 0 R /Size {len(objects) + 1} >>\nstartxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(content)


def _minimal_xlsx() -> bytes:
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types" />',
        )
        archive.writestr(
            "xl/workbook.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <sheets><sheet name="Sales" sheetId="1" r:id="rId1"/></sheets>
</workbook>""",
        )
        archive.writestr(
            "xl/_rels/workbook.xml.rels",
            """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Target="worksheets/sheet1.xml" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet"/>
</Relationships>""",
        )
        archive.writestr(
            "xl/sharedStrings.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<sst xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
  <si><t>Region</t></si><si><t>Revenue</t></si><si><t>East</t></si><si><t>10</t></si>
</sst>""",
        )
        archive.writestr(
            "xl/worksheets/sheet1.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
  <sheetData>
    <row r="1"><c r="A1" t="s"><v>0</v></c><c r="B1" t="s"><v>1</v></c></row>
    <row r="2"><c r="A2" t="s"><v>2</v></c><c r="B2" t="s"><v>3</v></c></row>
  </sheetData>
</worksheet>""",
        )
    return buffer.getvalue()


def _minimal_pptx() -> bytes:
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types" />',
        )
        archive.writestr(
            "ppt/presentation.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<p:presentation xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <p:sldIdLst><p:sldId id="256" r:id="rId1"/></p:sldIdLst>
</p:presentation>""",
        )
        archive.writestr(
            "ppt/_rels/presentation.xml.rels",
            """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Target="slides/slide1.xml" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide"/>
</Relationships>""",
        )
        archive.writestr(
            "ppt/slides/slide1.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<p:sld xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
  <p:cSld><p:spTree>
    <p:sp><p:txBody><a:p><a:r><a:t>Architecture</a:t></a:r></a:p></p:txBody></p:sp>
    <p:sp><p:txBody><a:p><a:r><a:t>Hybrid retrieval</a:t></a:r></a:p></p:txBody></p:sp>
  </p:spTree></p:cSld>
</p:sld>""",
        )
    return buffer.getvalue()


def _worker_config(batch_size: int = 1) -> WorkerConfig:
    return WorkerConfig(
        poll_interval_seconds=0,
        batch_size=batch_size,
        lease_duration=timedelta(minutes=5),
        lease_renew_interval_seconds=60,
        shutdown_grace_seconds=30,
        enabled_job_types=frozenset({"document_ingest"}),
        worker_instance_id="worker-1",
    )
