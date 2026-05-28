from __future__ import annotations

import zipfile
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from app.core.config import get_settings
from app.core.errors import UnsafeFileRejected, UnsupportedMediaType
from app.core.job_utils import redact_error_message
from app.ingest.chunking import ChunkingConfig, ChunkingError, FixedTokenChunker
from app.ingest.extractors.base import ExtractedDocument, ExtractedPage, ExtractionError
from app.ingest.extractors.csv import CsvExtractor
from app.ingest.extractors.dispatcher import ExtractorDispatcher
from app.ingest.extractors.docx import DocxExtractor
from app.ingest.extractors.markdown import MarkdownExtractor
from app.ingest.extractors.office import ExcelExtractor, PowerPointExtractor
from app.ingest.extractors.pdf import PdfTextExtractor
from app.ingest.extractors.text import PlainTextExtractor
from app.ingest.extractors.web import HtmlExtractor, XmlExtractor
from app.ingest.hashing import chunk_hash, normalize_chunk_text
from app.ingest.metadata import metadata_from_extracted_document
from app.storage import extractors as legacy_storage_extractors
from app.storage.validators import validate_upload


def test_text_extractor_decodes_utf8_utf8_sig_and_cp932(tmp_path: Path) -> None:
    extractor = PlainTextExtractor()
    cases = [
        ("utf8.txt", b"hello utf8", "hello utf8"),
        ("bom.txt", "\ufeffhello bom".encode(), "hello bom"),
        ("cp932.txt", b"\x82\xa0", "\u3042"),
    ]

    for file_name, content, expected in cases:
        path = tmp_path / file_name
        path.write_bytes(content)
        extracted = extractor.extract(path, _metadata(file_name, "text/plain", len(content)))
        assert extracted.pages[0].text == expected
        assert extracted.metadata.extractor_name == "plain_text"


def test_text_extractor_failure_and_empty_text(tmp_path: Path) -> None:
    extractor = PlainTextExtractor()
    invalid = tmp_path / "invalid.txt"
    invalid.write_bytes(b"\x81")
    with pytest.raises(ExtractionError) as invalid_exc:
        extractor.extract(invalid, _metadata("invalid.txt", "text/plain", 1))
    assert invalid_exc.value.error_code == "text_extraction_failed"

    empty = tmp_path / "empty.txt"
    empty.write_text("   \n\t", encoding="utf-8")
    with pytest.raises(ExtractionError) as empty_exc:
        extractor.extract(empty, _metadata("empty.txt", "text/plain", 5))
    assert empty_exc.value.error_code == "empty_extracted_text"


def test_markdown_extraction_keeps_atx_section_titles(tmp_path: Path) -> None:
    path = tmp_path / "guide.md"
    path.write_text("# Intro\nbody\n## Detail\nmore\n", encoding="utf-8")

    extracted = MarkdownExtractor().extract(path, _metadata("guide.md", "text/markdown", 27))

    assert [page.section_title for page in extracted.pages] == ["Intro", "Detail"]
    assert "# Intro" in extracted.pages[0].text
    assert extracted.metadata.extractor_name == "markdown"


def test_csv_extraction_and_malformed_handling(tmp_path: Path) -> None:
    path = tmp_path / "data.csv"
    path.write_text("name,value\nalpha,1\n", encoding="utf-8")

    extracted = CsvExtractor().extract(path, _metadata("data.csv", "text/csv", 19))

    assert extracted.pages[0].text == "name | value\nalpha | 1"

    bom_path = tmp_path / "bom.csv"
    bom_path.write_bytes("\ufeffname,value\nalpha,1\n".encode())
    bom_extracted = CsvExtractor().extract(bom_path, _metadata("bom.csv", "text/csv", 22))
    assert bom_extracted.pages[0].text.startswith("name | value")

    malformed = tmp_path / "bad.csv"
    malformed.write_text('"unterminated', encoding="utf-8")
    with pytest.raises(ExtractionError) as exc:
        CsvExtractor().extract(malformed, _metadata("bad.csv", "text/csv", 13))
    assert exc.value.error_code == "text_extraction_failed"


def test_docx_extraction_includes_paragraphs_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "doc.docx"
    document = Document()
    document.add_paragraph("Paragraph text")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    document.save(str(path))

    extracted = DocxExtractor().extract(
        path,
        _metadata(
            "doc.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            path.stat().st_size,
        ),
    )

    assert "Paragraph text" in extracted.pages[0].text
    assert "A | B" in extracted.pages[0].text
    assert extracted.metadata.extractor_name == "docx"


def test_docx_upload_validation_rejects_zip_bomb_shape() -> None:
    content = _docx_zip_with_large_document_xml()

    with pytest.raises(UnsafeFileRejected):
        validate_upload(
            filename="bomb.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            content=content,
            max_bytes=len(content) + 1,
            allowed_extensions=[".docx"],
        )


def test_pdf_text_layer_extraction(tmp_path: Path) -> None:
    path = tmp_path / "sample.pdf"
    path.write_bytes(_minimal_pdf("Hello PDF text"))

    extracted = PdfTextExtractor().extract(
        path, _metadata("sample.pdf", "application/pdf", path.stat().st_size)
    )

    assert "Hello PDF text" in extracted.pages[0].text
    assert extracted.pages[0].page_number == 1
    assert extracted.metadata.page_count == 1


def test_extractor_dispatcher_validates_extension_and_mime(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv(
        "UPLOAD_ALLOWED_EXTENSIONS",
        '[".pdf",".docx",".txt",".md",".markdown",".csv",".xlsx",".pptx",".html",".htm",".xml"]',
    )
    get_settings.cache_clear()
    dispatcher = ExtractorDispatcher()

    assert (
        dispatcher.select(file_name="a.pdf", mime_type="application/pdf").name == "pdf_text_layer"
    )
    assert dispatcher.select(file_name="a.markdown", mime_type="text/markdown").name == "markdown"
    assert (
        dispatcher.select(
            file_name="a.xlsx",
            mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ).name
        == "xlsx"
    )
    assert (
        dispatcher.select(
            file_name="a.pptx",
            mime_type="application/vnd.openxmlformats-officedocument.presentationml.presentation",
        ).name
        == "pptx"
    )
    assert dispatcher.select(file_name="a.html", mime_type="text/html").name == "html"
    assert dispatcher.select(file_name="a.xml", mime_type="application/xml").name == "xml"

    with pytest.raises(ExtractionError) as unsupported:
        dispatcher.select(file_name="a.exe", mime_type="application/octet-stream")
    assert unsupported.value.error_code == "unsupported_file_type"

    with pytest.raises(ExtractionError) as mismatch:
        dispatcher.select(file_name="a.pdf", mime_type="text/plain")
    assert mismatch.value.error_code == "mime_type_mismatch"
    get_settings.cache_clear()


def test_metadata_and_chunking_are_deterministic() -> None:
    document = ExtractedDocument(
        pages=[
            ExtractedPage("alpha beta gamma delta", page_number=1, section_title="Intro"),
            ExtractedPage("epsilon zeta eta theta", page_number=2, section_title="Next"),
        ],
        metadata=_extraction_metadata(page_count=2),
    )
    metadata = metadata_from_extracted_document(document)
    chunker = FixedTokenChunker(ChunkingConfig(chunk_size_tokens=5, chunk_overlap_tokens=2))

    first = chunker.chunk(document, document_version_id=10)
    second = chunker.chunk(document, document_version_id=10)

    assert first == second
    assert metadata.page_count == 2
    assert [chunk.chunk_index for chunk in first] == [0, 1, 2]
    assert first[0].token_count == 5
    assert first[0].char_count == len(first[0].content_text)
    assert first[0].page_from == 1
    assert first[0].page_to == 2
    assert first[0].section_title == "Intro"
    assert first[0].chunk_hash == chunk_hash(
        normalized_chunk_text=normalize_chunk_text(first[0].content_text),
        document_version_id=10,
        chunk_index=0,
    )


def test_xlsx_upload_validation_and_extraction_metadata(tmp_path: Path) -> None:
    content = _minimal_xlsx()
    validate_upload(
        filename="book.xlsx",
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        content=content,
        max_bytes=len(content) + 1,
        allowed_extensions=[".xlsx"],
    )
    path = tmp_path / "book.xlsx"
    path.write_bytes(content)

    extracted = ExcelExtractor().extract(
        path,
        _metadata(
            "book.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            len(content),
        ),
    )

    assert extracted.metadata.extractor_name == "xlsx"
    assert len(extracted.pages) == 2
    page = extracted.pages[0]
    assert "Sheet: Sales" in page.text
    assert "A=Region | B=Revenue" in page.text
    assert page.section_title == "Sheet: Sales"
    assert page.metadata["structure_type"] == "excel_sheet"
    assert page.metadata["sheet_name"] == "Sales"
    assert page.metadata["row_from"] == 1
    assert page.metadata["row_to"] == 2
    assert extracted.pages[1].metadata["sheet_name"] == "Hidden"
    assert "Hidden Value" in extracted.pages[1].text


def test_html_upload_validation_and_extraction_removes_active_content(tmp_path: Path) -> None:
    content = (
        b"<!doctype html><html><head><title>Roadmap</title><style>.x{}</style>"
        b"<script>secret()</script></head><body><h1>Phase 2</h1>"
        b"<p>HTML ingest text</p><table><tr><th>Name</th><td>Value</td></tr></table>"
        b"</body></html>"
    )
    validate_upload(
        filename="page.html",
        content_type="text/html",
        content=content,
        max_bytes=len(content) + 1,
        allowed_extensions=[".html"],
    )
    path = tmp_path / "page.html"
    path.write_bytes(content)

    extracted = HtmlExtractor().extract(path, _metadata("page.html", "text/html", len(content)))

    joined = "\n".join(page.text for page in extracted.pages)
    assert "Phase 2" in joined
    assert "HTML ingest text" in joined
    assert "Name | Value" in joined
    assert "secret()" not in joined
    assert extracted.pages[0].metadata["structure_type"] == "html_section"
    assert extracted.pages[0].metadata["parent_child_schema_version"] == "phase2.web_ingest.v1"


def test_web_extraction_preserves_long_body_text(tmp_path: Path) -> None:
    long_text = "x" * 1200
    html_content = f"<html><body><p>{long_text}</p></body></html>".encode()
    html_path = tmp_path / "long.html"
    html_path.write_bytes(html_content)

    html_extracted = HtmlExtractor().extract(
        html_path,
        _metadata("long.html", "text/html", len(html_content)),
    )

    assert html_extracted.pages[0].text == long_text

    xml_content = f"<root><entry>{long_text}</entry></root>".encode()
    xml_path = tmp_path / "long.xml"
    xml_path.write_bytes(xml_content)

    xml_extracted = XmlExtractor().extract(
        xml_path,
        _metadata("long.xml", "application/xml", len(xml_content)),
    )

    assert xml_extracted.pages[0].text == long_text


def test_xml_upload_validation_and_extraction_metadata(tmp_path: Path) -> None:
    content = b"""<?xml version="1.0" encoding="UTF-8"?>
<feed><entry><title>Release</title><summary>XML ingest text</summary></entry></feed>"""
    validate_upload(
        filename="feed.xml",
        content_type="application/xml",
        content=content,
        max_bytes=len(content) + 1,
        allowed_extensions=[".xml"],
    )
    path = tmp_path / "feed.xml"
    path.write_bytes(content)

    extracted = XmlExtractor().extract(path, _metadata("feed.xml", "application/xml", len(content)))

    assert extracted.metadata.extractor_name == "xml"
    assert any("XML ingest text" in page.text for page in extracted.pages)
    first = extracted.pages[0]
    assert first.metadata["structure_type"] == "xml_element"
    assert "feed" in str(first.metadata["xml_path"])


def test_xml_extraction_uses_parent_direct_text_without_child_duplicates(tmp_path: Path) -> None:
    content = b"<root>intro<child>A</child>middle<child>B</child>end</root>"
    path = tmp_path / "nested.xml"
    path.write_bytes(content)

    extracted = XmlExtractor().extract(
        path, _metadata("nested.xml", "application/xml", len(content))
    )

    texts = [page.text for page in extracted.pages]
    assert texts == ["intro middle end", "A", "B"]
    assert "intro A middle B end" not in texts


def test_xml_extraction_parent_keys_include_element_occurrence(tmp_path: Path) -> None:
    content = b"<feed><entry><title>A</title></entry><entry><title>B</title></entry></feed>"
    path = tmp_path / "feed.xml"
    path.write_bytes(content)

    extracted = XmlExtractor().extract(path, _metadata("feed.xml", "application/xml", len(content)))

    title_keys = [
        page.metadata["parent_chunk_key"]
        for page in extracted.pages
        if page.metadata["xml_path"] == "feed / entry / title"
    ]
    assert len(title_keys) == 2
    assert len(set(title_keys)) == 2


def test_xml_extraction_enforces_element_limit(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("INGEST_XML_MAX_ELEMENTS", "2")
    get_settings.cache_clear()
    content = b"<root><child>A</child><child>B</child></root>"
    path = tmp_path / "too-many.xml"
    path.write_bytes(content)

    with pytest.raises(ExtractionError):
        XmlExtractor().extract(path, _metadata("too-many.xml", "application/xml", len(content)))

    get_settings.cache_clear()


def test_xml_entities_and_svg_are_rejected(tmp_path: Path) -> None:
    with pytest.raises(UnsafeFileRejected):
        validate_upload(
            filename="page.html",
            content_type="application/xhtml+xml",
            content=b"""<?xml version="1.0"?><!DOCTYPE html [<!ENTITY unsafe "x">]><html />""",
            max_bytes=1000,
            allowed_extensions=[".html"],
        )

    with pytest.raises(UnsafeFileRejected):
        validate_upload(
            filename="vector.xml",
            content_type="application/xml",
            content=b"<svg><script>alert(1)</script></svg>",
            max_bytes=1000,
            allowed_extensions=[".xml"],
        )

    with pytest.raises(UnsafeFileRejected):
        validate_upload(
            filename="prefixed-vector.xml",
            content_type="application/xml",
            content=b'<x:svg xmlns:x="http://www.w3.org/2000/svg"><x:text>unsafe</x:text></x:svg>',
            max_bytes=1000,
            allowed_extensions=[".xml"],
        )

    svg_path = tmp_path / "prefixed-vector.xml"
    svg_path.write_text(
        '<x:svg xmlns:x="http://www.w3.org/2000/svg"><x:text>unsafe</x:text></x:svg>',
        encoding="utf-8",
    )
    with pytest.raises(ExtractionError):
        XmlExtractor().extract(
            svg_path, _metadata("prefixed-vector.xml", "application/xml", svg_path.stat().st_size)
        )

    path = tmp_path / "entity.xml"
    path.write_text(
        """<?xml version="1.0"?><!DOCTYPE x [<!ENTITY unsafe "expanded">]><x>&unsafe;</x>""",
        encoding="utf-8",
    )
    with pytest.raises(ExtractionError) as exc:
        XmlExtractor().extract(
            path, _metadata("entity.xml", "application/xml", path.stat().st_size)
        )
    assert exc.value.error_code == "text_extraction_failed"

    with pytest.raises(UnsafeFileRejected):
        validate_upload(
            filename="late-entity.xml",
            content_type="application/xml",
            content=(b"<root>" + (b"x" * 5000) + b"<!ENTITY unsafe 'expanded'></root>"),
            max_bytes=6000,
            allowed_extensions=[".xml"],
        )


def test_xlsx_extraction_skips_hidden_sheets(tmp_path: Path) -> None:
    content = _minimal_xlsx(hidden_second_sheet=True)
    path = tmp_path / "hidden.xlsx"
    path.write_bytes(content)

    extracted = ExcelExtractor().extract(
        path,
        _metadata(
            "hidden.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            len(content),
        ),
    )

    assert len(extracted.pages) == 1
    assert "Hidden" not in extracted.pages[0].text


def test_xlsx_extraction_uses_formula_text_when_cached_value_missing(
    tmp_path: Path,
) -> None:
    content = _minimal_xlsx(
        sheet1_xml="""<?xml version="1.0" encoding="UTF-8"?>
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
  <sheetData>
    <row r="1"><c r="A1"><f>SUM(B1:C1)</f></c></row>
  </sheetData>
</worksheet>""",
    )
    path = tmp_path / "formula.xlsx"
    path.write_bytes(content)

    extracted = ExcelExtractor().extract(
        path,
        _metadata(
            "formula.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            len(content),
        ),
    )

    assert len(extracted.pages) == 2
    assert "A=SUM(B1:C1)" in extracted.pages[0].text


def test_xlsx_extraction_rejects_xml_entities(tmp_path: Path) -> None:
    content = _minimal_xlsx(
        workbook_xml="""<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE workbook [<!ENTITY unsafe "expanded">]>
<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main"
  xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <sheets><sheet name="Sales" sheetId="1" r:id="rId1"/></sheets>
</workbook>""",
    )
    path = tmp_path / "entities.xlsx"
    path.write_bytes(content)

    with pytest.raises(ExtractionError) as exc:
        ExcelExtractor().extract(
            path,
            _metadata(
                "entities.xlsx",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                len(content),
            ),
        )
    assert exc.value.error_code == "text_extraction_failed"


def test_pptx_upload_validation_and_extraction_metadata(tmp_path: Path) -> None:
    content = _minimal_pptx()
    validate_upload(
        filename="deck.pptx",
        content_type="application/vnd.openxmlformats-officedocument.presentationml.presentation",
        content=content,
        max_bytes=len(content) + 1,
        allowed_extensions=[".pptx"],
    )
    path = tmp_path / "deck.pptx"
    path.write_bytes(content)

    extracted = PowerPointExtractor().extract(
        path,
        _metadata(
            "deck.pptx",
            "application/vnd.openxmlformats-officedocument.presentationml.presentation",
            len(content),
        ),
    )

    assert extracted.metadata.extractor_name == "pptx"
    assert extracted.metadata.page_count == 1
    page = extracted.pages[0]
    assert "Slide 1: Architecture" in page.text
    assert "Shape 2: Hybrid retrieval" in page.text
    assert page.page_number == 1
    assert page.metadata["structure_type"] == "powerpoint_slide"
    assert page.metadata["slide_number"] == 1


def test_legacy_storage_extractor_supports_office_text(tmp_path: Path) -> None:
    xlsx_path = tmp_path / "book.xlsx"
    pptx_path = tmp_path / "deck.pptx"
    xlsx_path.write_bytes(_minimal_xlsx())
    pptx_path.write_bytes(_minimal_pptx())

    assert "Sheet: Sales" in legacy_storage_extractors.extract_text(xlsx_path)
    assert "Slide 1: Architecture" in legacy_storage_extractors.extract_text(pptx_path)


def test_macro_enabled_office_files_are_rejected() -> None:
    with pytest.raises(UnsafeFileRejected):
        validate_upload(
            filename="macro.xlsm",
            content_type="application/vnd.ms-excel.sheet.macroEnabled.12",
            content=_minimal_xlsx(),
            max_bytes=100000,
            allowed_extensions=[".xlsx", ".xlsm"],
        )


@pytest.mark.parametrize(
    ("filename", "content_type"),
    [
        ("legacy.xls", "application/vnd.ms-excel"),
        ("legacy.ppt", "application/vnd.ms-powerpoint"),
    ],
)
def test_legacy_office_formats_are_not_supported(filename: str, content_type: str) -> None:
    with pytest.raises(UnsupportedMediaType):
        validate_upload(
            filename=filename,
            content_type=content_type,
            content=b"legacy office content",
            max_bytes=100000,
            allowed_extensions=[".xlsx", ".pptx"],
        )


@pytest.mark.parametrize(
    "part_name",
    [
        "xl/vbaProject.bin",
        "xl/VBAPROJECT.BIN",
        "xl/embeddings/oleObject1.bin",
        "xl/Embeddings/oleObject1.bin",
    ],
)
def test_office_archives_with_macro_or_embedded_parts_are_rejected(part_name: str) -> None:
    with pytest.raises(UnsafeFileRejected):
        validate_upload(
            filename="book.xlsx",
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            content=_minimal_xlsx(extra_entries={part_name: b"unsafe"}),
            max_bytes=100000,
            allowed_extensions=[".xlsx"],
        )


def test_parent_child_metadata_is_preserved_on_chunks() -> None:
    document = ExtractedDocument(
        pages=[
            ExtractedPage(
                "Region Revenue East 10",
                section_title="Sheet: Sales",
                metadata={
                    "parent_child_schema_version": "phase2.parent_child.v1",
                    "structure_type": "excel_sheet",
                    "parent_chunk_key": "xlsx:sheet:1",
                    "parent_title": "Sales",
                    "sheet_name": "Sales",
                    "row_from": 1,
                    "row_to": 2,
                    "column_from": 1,
                    "column_to": 2,
                    "table_index": 1,
                },
            )
        ],
        metadata=_extraction_metadata(page_count=None),
    )

    chunks = FixedTokenChunker(ChunkingConfig(chunk_size_tokens=4, chunk_overlap_tokens=0)).chunk(
        document, document_version_id=12
    )

    assert chunks[0].metadata_json is not None
    assert chunks[0].metadata_json["parent_chunk_key"] == "xlsx:sheet:1"
    assert chunks[0].metadata_json["child_chunk_key"] == "xlsx:sheet:1:chunk:0"
    assert chunks[0].metadata_json["chunk_level"] == "child"


def test_parent_child_chunking_does_not_cross_parent_boundaries() -> None:
    document = ExtractedDocument(
        pages=[
            ExtractedPage(
                "Region Revenue East 10",
                section_title="Sheet: Sales",
                metadata={
                    "parent_child_schema_version": "phase2.parent_child.v1",
                    "structure_type": "excel_sheet",
                    "parent_chunk_key": "xlsx:sheet:1",
                    "parent_title": "Sales",
                    "sheet_name": "Sales",
                    "row_from": 1,
                    "row_to": 2,
                    "column_from": 1,
                    "column_to": 2,
                    "table_index": 1,
                },
            ),
            ExtractedPage(
                "Owner Budget West 20",
                section_title="Sheet: Budget",
                metadata={
                    "parent_child_schema_version": "phase2.parent_child.v1",
                    "structure_type": "excel_sheet",
                    "parent_chunk_key": "xlsx:sheet:2",
                    "parent_title": "Budget",
                    "sheet_name": "Budget",
                    "row_from": 1,
                    "row_to": 2,
                    "column_from": 1,
                    "column_to": 2,
                    "table_index": 1,
                },
            ),
        ],
        metadata=_extraction_metadata(page_count=None),
    )

    chunks = FixedTokenChunker(ChunkingConfig(chunk_size_tokens=20, chunk_overlap_tokens=0)).chunk(
        document, document_version_id=12
    )

    assert len(chunks) == 2
    assert [chunk.section_title for chunk in chunks] == ["Sheet: Sales", "Sheet: Budget"]
    metadata_keys = []
    for chunk in chunks:
        assert chunk.metadata_json is not None
        metadata_keys.append(chunk.metadata_json["parent_chunk_key"])
    assert metadata_keys == ["xlsx:sheet:1", "xlsx:sheet:2"]


def test_chunking_rejects_invalid_config_and_no_chunks() -> None:
    with pytest.raises(ValueError):
        ChunkingConfig(chunk_size_tokens=5, chunk_overlap_tokens=5)

    empty = ExtractedDocument(
        pages=[ExtractedPage("   ", page_number=None)],
        metadata=_extraction_metadata(page_count=None),
    )
    with pytest.raises(ChunkingError) as exc:
        FixedTokenChunker(ChunkingConfig(chunk_size_tokens=5, chunk_overlap_tokens=1)).chunk(
            empty, document_version_id=1
        )
    assert exc.value.error_code == "no_chunks_created"


def test_chunking_overlap_zero_and_size_boundaries() -> None:
    exact = ExtractedDocument(
        pages=[ExtractedPage("one two three", page_number=1)],
        metadata=_extraction_metadata(page_count=1),
    )
    plus_one = ExtractedDocument(
        pages=[ExtractedPage("one two three four", page_number=1)],
        metadata=_extraction_metadata(page_count=1),
    )
    chunker = FixedTokenChunker(ChunkingConfig(chunk_size_tokens=3, chunk_overlap_tokens=0))

    exact_chunks = chunker.chunk(exact, document_version_id=1)
    plus_one_chunks = chunker.chunk(plus_one, document_version_id=1)

    assert [chunk.content_text for chunk in exact_chunks] == ["one two three"]
    assert [chunk.content_text for chunk in plus_one_chunks] == ["one two three", "four"]


def test_redaction_helper_removes_sensitive_error_content() -> None:
    assert redact_error_message("chunk content_text contained raw document text") == (
        "Job failed with a redacted error."
    )


def _metadata(file_name: str, mime_type: str, file_size_bytes: int):
    from app.ingest.extractors.base import ExtractionInputMetadata

    return ExtractionInputMetadata(
        file_name=file_name,
        mime_type=mime_type,
        file_size_bytes=file_size_bytes,
    )


def _extraction_metadata(page_count: int | None):
    from app.ingest.extractors.base import ExtractionMetadata

    return ExtractionMetadata(
        extractor_name="test",
        extractor_version="1",
        page_count=page_count,
    )


def _minimal_pdf(text: str) -> bytes:
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("ascii")
    objects = [
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n",
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n",
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >> endobj\n",
        b"4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n",
        b"5 0 obj << /Length "
        + str(len(stream)).encode("ascii")
        + b" >> stream\n"
        + stream
        + b"\nendstream endobj\n",
    ]
    content = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for obj in objects:
        offsets.append(len(content))
        content.extend(obj)
    xref_offset = len(content)
    content.extend(f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode("ascii"))
    for offset in offsets[1:]:
        content.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    content.extend(
        (
            f"trailer << /Root 1 0 R /Size {len(objects) + 1} >>\nstartxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(content)


def _docx_zip_with_large_document_xml() -> bytes:
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "<Types />")
        archive.writestr("word/document.xml", "x" * (5 * 1024 * 1024 + 1))
    return buffer.getvalue()


def _minimal_xlsx(
    *,
    hidden_second_sheet: bool = False,
    workbook_xml: str | None = None,
    sheet1_xml: str | None = None,
    extra_entries: dict[str, bytes] | None = None,
) -> bytes:
    buffer = BytesIO()
    second_sheet_state = ' state="hidden"' if hidden_second_sheet else ""
    with zipfile.ZipFile(buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types" />',
        )
        archive.writestr(
            "xl/workbook.xml",
            workbook_xml
            or f"""<?xml version="1.0" encoding="UTF-8"?>
<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <sheets>
    <sheet name="Sales" sheetId="1" r:id="rId1"/>
    <sheet name="Hidden" sheetId="2" r:id="rId2"{second_sheet_state}/>
  </sheets>
</workbook>""",
        )
        archive.writestr(
            "xl/_rels/workbook.xml.rels",
            """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Target="worksheets/sheet1.xml" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet"/>
  <Relationship Id="rId2" Target="worksheets/sheet2.xml" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet"/>
</Relationships>""",
        )
        archive.writestr(
            "xl/sharedStrings.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<sst xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
  <si><t>Region</t></si>
  <si><t>Revenue</t></si>
  <si><t>East</t></si>
  <si><t>10</t></si>
  <si><t>Hidden Value</t></si>
</sst>""",
        )
        archive.writestr(
            "xl/worksheets/sheet1.xml",
            sheet1_xml
            or """<?xml version="1.0" encoding="UTF-8"?>
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
  <sheetData>
    <row r="1"><c r="A1" t="s"><v>0</v></c><c r="B1" t="s"><v>1</v></c></row>
    <row r="2"><c r="A2" t="s"><v>2</v></c><c r="B2" t="s"><v>3</v></c></row>
  </sheetData>
</worksheet>""",
        )
        archive.writestr(
            "xl/worksheets/sheet2.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
  <sheetData>
    <row r="1"><c r="A1" t="s"><v>4</v></c></row>
  </sheetData>
</worksheet>""",
        )
        for name, payload in (extra_entries or {}).items():
            archive.writestr(name, payload)
    return buffer.getvalue()


def _minimal_pptx() -> bytes:
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types" />',
        )
        archive.writestr(
            "ppt/presentation.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<p:presentation xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <p:sldIdLst><p:sldId id="256" r:id="rId1"/></p:sldIdLst>
</p:presentation>""",
        )
        archive.writestr(
            "ppt/_rels/presentation.xml.rels",
            """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Target="slides/slide1.xml" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide"/>
</Relationships>""",
        )
        archive.writestr(
            "ppt/slides/slide1.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<p:sld xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
  <p:cSld><p:spTree>
    <p:sp><p:txBody><a:p><a:r><a:t>Architecture</a:t></a:r></a:p></p:txBody></p:sp>
    <p:sp><p:txBody><a:p><a:r><a:t>Hybrid retrieval</a:t></a:r></a:p></p:txBody></p:sp>
  </p:spTree></p:cSld>
</p:sld>""",
        )
    return buffer.getvalue()
