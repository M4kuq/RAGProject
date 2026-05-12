from __future__ import annotations

from pathlib import Path

from docx import Document

from app.ingest.extractors.base import (
    ExtractedDocument,
    ExtractedPage,
    ExtractionInputMetadata,
    ExtractionMetadata,
    ensure_non_empty_text,
    safe_extraction_failure,
)

DOCX_EXTRACTOR_VERSION = "1"


class DocxExtractor:
    name = "docx"
    version = DOCX_EXTRACTOR_VERSION

    def extract(self, file_path: Path, metadata: ExtractionInputMetadata) -> ExtractedDocument:
        try:
            document = Document(str(file_path))
            lines = [paragraph.text.strip() for paragraph in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        lines.append(" | ".join(cells))
        except Exception as exc:
            raise safe_extraction_failure() from exc

        text = "\n".join(line for line in lines if line)
        pages = [ExtractedPage(text=text, page_number=None)]
        ensure_non_empty_text(pages)
        return ExtractedDocument(
            pages=pages,
            metadata=ExtractionMetadata(
                extractor_name=self.name,
                extractor_version=self.version,
                page_count=None,
            ),
        )
